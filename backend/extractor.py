"""
Pulls plain text out of an uploaded old resume (.docx, .pdf, or .txt).

Works on in-memory file-like objects (BytesIO from an upload) rather than
disk paths, since this runs inside a web request.
"""
import io


def extract_docx_text(file_obj: io.BytesIO) -> str:
    from docx import Document

    doc = Document(file_obj)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return "\n".join(lines)


def extract_pdf_text(file_obj: io.BytesIO) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_obj)
    lines = []
    for page in reader.pages:
        text = page.extract_text() or ""
        for line in text.splitlines():
            line = line.strip()
            if line:
                lines.append(line)
    return "\n".join(lines)


def extract_text(filename: str, content: bytes) -> str:
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    file_obj = io.BytesIO(content)

    if suffix == "docx":
        return extract_docx_text(file_obj)
    elif suffix == "pdf":
        return extract_pdf_text(file_obj)
    elif suffix == "txt":
        return content.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: .{suffix} (expected .docx, .pdf, or .txt)")
