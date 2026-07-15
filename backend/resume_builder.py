"""
Builds a .docx resume matching the WrenPath template layout, in memory
(returns bytes rather than writing to disk, since this runs inside a web
request):

  NAME (bold, centered)
  Location | Phone | Email | LinkedIn   (centered)
  [shaded bar] PROFESSIONAL SUMMARY
  Summary paragraph(s)
  [shaded bar] CORE SKILLS & EXPERTISE
  Skills bullets, two columns by default (or one in ATS mode)
  [shaded bar] PROFESSIONAL EXPERIENCE
  Job title (bold) / Company, Location - Dates
  Bulleted achievements
  ... (repeat per job)
  [shaded bar] EDUCATION
  Bulleted degree/school lines
"""
import io

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADING_FILL = "CADCF2"
FONT_NAME = "Calibri"
FONT_SIZE = Pt(11)
PAGE_MARGIN = Inches(0.7)


def shade_paragraph(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def set_columns(section, num_cols):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), "720")


def style_run(run, bold=False):
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_heading_bar(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    style_run(run, bold=True)
    shade_paragraph(p, HEADING_FILL)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    style_run(run)
    return p


def set_page_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = PAGE_MARGIN
    section.right_margin = PAGE_MARGIN
    section.top_margin = PAGE_MARGIN
    section.bottom_margin = PAGE_MARGIN


def build_resume_bytes(data: dict, ats_mode: bool = False) -> bytes:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE

    set_page_geometry(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = p.add_run(data["name"])
    style_run(name_run, bold=True)
    p.add_run().add_break()
    contact_run = p.add_run(data["contact"])
    style_run(contact_run)

    add_heading_bar(doc, "PROFESSIONAL SUMMARY")

    summary = data["summary"]
    paragraphs = summary if isinstance(summary, list) else [summary]
    p = doc.add_paragraph()
    for i, para_text in enumerate(paragraphs):
        if i > 0:
            p.add_run().add_break()
            p.add_run().add_break()
        run = p.add_run(para_text)
        style_run(run)

    # Skills, experience, and education sections only render if there's
    # actual content - an empty list used to still print the heading bar
    # with nothing underneath it (e.g. someone using the start-from-scratch
    # wizard who added zero experience entries got an orphaned
    # "PROFESSIONAL EXPERIENCE" heading followed immediately by EDUCATION).
    if data["skills"]:
        doc.add_section(WD_SECTION.CONTINUOUS)
        set_page_geometry(doc.sections[-1])
        add_heading_bar(doc, "CORE SKILLS & EXPERTISE")

        doc.add_section(WD_SECTION.CONTINUOUS)
        set_page_geometry(doc.sections[-1])
        set_columns(doc.sections[-1], 1 if ats_mode else 2)
        for skill in data["skills"]:
            add_bullet(doc, skill)

    # Always start a fresh single-column section here, regardless of
    # whether there's experience content - this undoes the skills
    # section's column count (if it was 2), and EDUCATION below always
    # needs single column too.
    doc.add_section(WD_SECTION.CONTINUOUS)
    set_page_geometry(doc.sections[-1])
    set_columns(doc.sections[-1], 1)

    if data["experience"]:
        add_heading_bar(doc, "PROFESSIONAL EXPERIENCE")
        for i, job in enumerate(data["experience"]):
            p = doc.add_paragraph()
            if i > 0:
                p.paragraph_format.space_before = Pt(12)
            title_run = p.add_run(job["title"])
            style_run(title_run, bold=True)
            p.add_run().add_break()
            subtitle_run = p.add_run(job["subtitle"])
            style_run(subtitle_run)
            for bullet in job["bullets"]:
                add_bullet(doc, bullet)

    if data["education"]:
        edu_heading = add_heading_bar(doc, "EDUCATION")
        edu_heading.paragraph_format.space_before = Pt(6)
        for entry in data["education"]:
            add_bullet(doc, entry)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
